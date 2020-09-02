import ast
import json
import os
import shutil
import urllib
import time
from datetime import datetime, timedelta
from io import BytesIO

import psycopg2
import requests
import geopandas as gpd
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from shapely.geometry import Polygon

def bbox_image(lat, lng, zoom):
    # produce bounding box of area they're viewing (fix to GR if zoomed out too far)
    if zoom < 11:
        zoom = 11
        lat = 42.965
        lng = -85.6719

    url = f'https://api.mapbox.com/styles/v1/mapbox/streets-v11/static/{lng},{lat},{zoom},0,0/1200x600?access_token=pk.eyJ1IjoiZnJlZXN0b2siLCJhIjoiY2p5bGd1dWg5MDhiZjNvdDgyYmFta2dweCJ9.5pV_bAq161_KlSJgGw0hYQ'
    response = requests.get(url)
    binary_img = BytesIO(response.content)
    return binary_img

# iterates through a table and changes the font size
def table_font_size(table,font_size):
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(font_size)
    return table

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    
# takes pandas dataframe and creates a docx table from it                    
def build_table(df,document):
    # build table
    table = document.add_table(rows=1,cols=df.shape[1])
    hdr_cells = table.rows[0].cells

    # create columns
    for n, i in enumerate(df.columns):
        hdr_cells[n].text = i

    # add df data to word table
    for col in range(df.shape[0]):
        row_cells = table.add_row().cells
        row = df.iloc[col]
        for n, i in enumerate(row):
            row_cells[n].text = str(i)

    # change font size of table
    table = table_font_size(table,font_size=9)
    
    return document

def df_to_word(title,df, map_image, graph, case_type, case_status):
    # create new document w/ narrow margins and headings
    document = Document()
    section = document.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    h = document.add_heading('',0)
    r = h.add_run()
    # r.add_picture(pic,width=Inches(1.0))
    r.add_text(title)
    
    num_cases = df['count'].sum()
    document.add_heading(f'Cases in Area: {num_cases}')
    document.add_picture(map_image, width=Inches(7))

    if not df.empty:
        document.add_picture(graph, width=Inches(7))
    else:
        return document # if empty, not dataframe information to add
    

    # add info about case statuses if relevant
    if case_type and not case_status:
        table = df[['status','count']].groupby(['status']).sum().reset_index()
        table.columns = ['Status','Count']
        document = build_table(table,document)
    elif case_status:
        pass
    else:
        table = df[['case_type','count']].groupby(['case_type']).sum().reset_index()

        table.columns = ['Case Type','Count']
        document = build_table(table,document)
    
    return document

def get_title(open_close):
    if open_close == True:
        title = 'Open Cases'
    elif open_close == False:
        title = 'Closed Cases'
    elif open_close == 'all':
        title = 'Open & Closed Cases'
    
    return title

def query_carto(query):
    ## CARTO ######################################
    url = 'https://freestone.carto.com/api/v2/sql'
    # query = "SELECT count(*) FROM cases_stripped"
    return requests.get(url,params={'q':query})
    
def report(uid, filters, bbox, lat, lng, zoom):
    map_image = bbox_image(lat, lng, zoom)

    # unpack filters
    open_close, case_type, case_status = filters
    if case_type:
        case_type = case_type.replace('"','')
    if case_status:
        case_status = case_status.replace('"','')


    # create query variables
    geom_query   = None
    open_query   = None
    status_query = None
    type_query   = None

    # construct polygon
    if zoom > 11:
        print('CREATE BOUNDING')
        # print(zoom)
        # print(bbox)
        x1, y1, x2, y2 = bbox
        bounding = Polygon([[x1,y1],[x1,y2],[x2,y2],[x2,y1]])
        # print(bounding)
        geom_query = "ST_Within(ST_AsText(the_geom), '{}')".format(bounding.wkt)
    else:
        print('NOT CREATING BOUNDING')
        # print(zoom)

    
    # establish queries based on user input
    if open_close == 'all':
        open_query = 'open IN (true, false)'
        print('ALL')
        if case_status:
            status_query = "status = '{}'".format(case_status)
            type_query = "case_type = '{}'".format(case_type)
            base_query = "SELECT count(date), right(date,4) dt, status  FROM cases_stripped"
            group_query = "dt, status"
        elif case_type:
            type_query = "case_type = '{}'".format(case_type)
            base_query = "SELECT count(date), right(date,4) dt, status  FROM cases_stripped"
            group_query = "dt, status"
        else:
            ## status or case type not specified
            base_query = "SELECT count(date), right(date,4) dt, case_type  FROM cases_stripped"
            group_query = "dt, case_type"
    else:
        print('Open_Close:',open_close)
        open_query = 'open = {}'.format(open_close)
        if case_status:
            status_query = "status = '{}'".format(case_status)
            type_query = "case_type = '{}'".format(case_type)
            base_query = "SELECT count(date), right(date,4) dt, status  FROM cases_stripped"
            group_query = "dt, status"
        elif case_type:
            print('Case Type:',case_type)
            type_query = "case_type = '{}'".format(case_type)
            base_query = "SELECT count(date), right(date,4) dt, status  FROM cases_stripped"
            group_query = "dt, status"
        else: 
            ## status or case type not specified
            base_query = "SELECT count(date), right(date,4) dt, case_type  FROM cases_stripped"
            group_query = "dt, case_type"

    query_strings = [geom_query, open_query, status_query, type_query]
    query_filter = [i for i in query_strings if i]
    query_insert = ' AND '.join(query_filter)
    query = "{} WHERE {} GROUP BY {}".format(base_query, query_insert, group_query)
    print(query)
    result = query_carto(query)

    # transform request to pandas dataframe
    title = get_title(open_close)
    table = result.json()['rows']
    
    if table:
        df = pd.DataFrame(table)
        print('-------------------------')
        print(df['count'].sum())
        print('------------------------')
        graph = create_graph(df, open_close, case_type, case_status)
        doc = df_to_word(f'Code Compliance - {title}', df, map_image, graph, case_type, case_status)
        f = BytesIO()
        doc.save(f)

        doc_binary = psycopg2.Binary(f.getvalue())
        update_data(uid, doc_binary)
    else:
        print('EMPTY DATAFRAME')
        graph = None 
        df = pd.DataFrame()
        doc = df_to_word(f'Code Compliance - {title}', df, map_image, graph, case_type, case_status)
        f = BytesIO()
        doc.save(f)
        doc_binary = psycopg2.Binary(f.getvalue())
        update_data(uid, doc_binary)


def create_graph(df, open_close, case_type, case_status):
    colors = ['skyblue','dodgerblue','slategrey','palegreen','mediumseagreen','lightcoral',
              'indianred','bisque','darkorange','thistle','mediumpurple','lightgray','peru',
              'lightseagreen','maroon']

    if not case_type:
        title = get_title(open_close)
        pivot = df.pivot_table(columns='case_type',index='dt',values='count')
        # cross = pd.crosstab(df.Date,df['Case Type'])
    elif not case_status:
        title = case_type
        pivot = df.pivot_table(columns='status',index='dt',values='count')
        # cross = pd.crosstab(df.Date,df.Status)
    else:
        title = f'{case_type} - {case_status}'
        # cross = pd.crosstab(df.Date,df.Status)
        pivot = df.pivot_table(columns='status',index='dt',values='count')

    bar = pivot.plot.bar(title=title,stacked=True,color=colors,figsize=(20,10))
    bar.legend(bbox_to_anchor=(1,1))
    bar.set_xlabel('Date')
    graph = BytesIO()
    bar.figure.savefig(graph,format='png',bbox_inches='tight')
    return graph

def fetch_data():
    conn = psycopg2.connect(os.environ['DATABASE_URL'],sslmode='require')
    cur = conn.cursor()
    
    try:
        cur.execute('SELECT * FROM reports')
        return cur.fetchall()
    except Exception as e:
        print('Error retrieving report parameters')
        print(e)
    finally:
        cur.close()
        conn.close()

def delete_data(uid):
    conn = psycopg2.connect(os.environ['DATABASE_URL'],sslmode='require')
    cur = conn.cursor()
    
    try:
        cur.execute('DELETE FROM reports WHERE uid = %s;',(uid,))
        conn.commit()
    except Exception as e:
        print('Error deleting report row')
        print(e)
    finally:
        cur.close()
        conn.close()

def update_data(uid, report_file):
    conn = psycopg2.connect(os.environ['DATABASE_URL'],sslmode='require')
    cur = conn.cursor()
    
    try:
        cur.execute('UPDATE reports SET completed = %s WHERE uid = %s;',(True,uid))
        cur.execute('UPDATE reports SET report_file = %s WHERE uid = %s;',(report_file,uid))
        conn.commit()
    except Exception as e:
        print('Error inserting report parameters')
        print(e)
    finally:
        cur.close()
        conn.close()

def initialize():
    while True:
        print('Checking...')
        parameters = fetch_data()
        for p in parameters:
            print('---------------------')
            uid     = p[0]
            completed  = p[7]
            created    = p[8]
            downloaded = p[9]
            if downloaded:
                print('Report downloaded, deleting record.')
                delete_data(uid)
                continue

            if completed and created < (datetime.now() - timedelta(hours = 1)):
                print('Deleting report that has been created but not downloaded within the hour.')
                delete_data(uid)
                continue

            if completed:
                print('Report already created, skipping')
                continue

            form    = p[1]
            lng     = p[2]
            lat     = p[3]
            zoom    = p[4]
            filters = p[5]
            bbox    = p[6]

            # yes, I'm not proud of how I stored and unpacked the filters in the database
            filters = filters.replace('{','').replace('}','').split(',')
            if filters[0] == 'Open':   filters[0] = True
            if filters[0] == 'Closed': filters[0] = False
            if filters[1] == 'NULL':   filters[1] = None
            if filters[2] == 'NULL':   filters[2] = None

            bbox = [float(i) for i in bbox.replace('{','').replace('}','').split(',')]

            if form == 'docx':
                report(uid, filters, bbox, lat, lng, zoom)
                
        time.sleep(10)

if __name__ == '__main__':
    initialize()



